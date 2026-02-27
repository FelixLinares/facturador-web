# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  Facturador FL v3.0  â€”  Backend unificado con PostgreSQL
#
#  DETECCIÃ“N AUTOMÃTICA:
#    - Si existe DATABASE_URL (Render) â†’ usa PostgreSQL
#    - Si no existe                    â†’ usa SQLite local
#
#  MÃ“DULOS:
#   /api/auth/*       â†’ login, logout, sesiÃ³n
#   /api/admin/*      â†’ gestiÃ³n de usuarios
#   /api/medical/*    â†’ facturador mÃ©dico + historial
#   /api/personal/*   â†’ facturador personal
#   /api/tasks/*      â†’ tareas y recordatorios por usuario
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

from pathlib import Path
from datetime import datetime, timedelta
from functools import wraps
import os, json, hashlib, secrets, uuid

from flask import Flask, jsonify, request, send_file, send_from_directory, g
from flask_cors import CORS

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas

# â”€â”€â”€ Base dirs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
BASE_DIR  = Path(__file__).parent
FRONT_DIR = BASE_DIR.parent / "frontend"
TEMP_DIR  = BASE_DIR / "temp"
TEMP_DIR.mkdir(exist_ok=True)

# â”€â”€â”€ Flask â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
app = Flask(
    __name__,
    template_folder=str(FRONT_DIR),
    static_folder=str(FRONT_DIR),
    static_url_path=""
)
CORS(app, supports_credentials=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  BASE DE DATOS  â€”  PostgreSQL en Render / SQLite local
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

DATABASE_URL = os.environ.get("DATABASE_URL", "")

if DATABASE_URL:
    # â”€â”€ PostgreSQL (Render) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    import psycopg2
    import psycopg2.extras

    # Render usa postgres:// pero psycopg2 necesita postgresql://
    if DATABASE_URL.startswith("postgres://"):
        DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

    def get_db():
        conn = psycopg2.connect(DATABASE_URL)
        conn.autocommit = True
        return conn

    def db_execute(sql, params=(), fetch="none"):
        conn = get_db()
        cur  = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params)
        if fetch == "one":  result = cur.fetchone()
        elif fetch == "all": result = cur.fetchall()
        else: result = None
        cur.close(); conn.close()
        return result

    PLACEHOLDER = "%s"
    print("âœ… Usando PostgreSQL (Render)")

else:
    # â”€â”€ SQLite (local) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    import sqlite3

    DB_PATH = BASE_DIR / "data" / "facturador.db"
    DB_PATH.parent.mkdir(exist_ok=True)

    def get_db():
        conn = sqlite3.connect(str(DB_PATH))
        conn.row_factory = sqlite3.Row
        return conn

    def db_execute(sql, params=(), fetch="none"):
        # Convierte %s â†’ ? para SQLite
        sql  = sql.replace("%s", "?")
        conn = get_db()
        cur  = conn.cursor()
        cur.execute(sql, params)
        if fetch == "one":  result = cur.fetchone()
        elif fetch == "all": result = cur.fetchall()
        else: result = None
        conn.commit(); conn.close()
        return result

    PLACEHOLDER = "?"
    print("âœ… Usando SQLite (local)")


def row_to_dict(row):
    if row is None: return None
    if isinstance(row, dict): return dict(row)
    return dict(row)

def rows_to_list(rows):
    if not rows: return []
    return [dict(r) for r in rows]

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  INICIALIZAR TABLAS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def init_db():
    if DATABASE_URL:
        # PostgreSQL â€” SERIAL para autoincrement
        db_execute("""
            CREATE TABLE IF NOT EXISTS users (
                id       TEXT PRIMARY KEY,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL,
                name     TEXT NOT NULL,
                role     TEXT DEFAULT 'user',
                active   BOOLEAN DEFAULT TRUE,
                modules  TEXT DEFAULT 'medical',
                created  TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS sessions (
                token   TEXT PRIMARY KEY,
                user_id TEXT NOT NULL,
                expires TEXT NOT NULL
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS medical_history (
                id             TEXT PRIMARY KEY,
                owner          TEXT NOT NULL,
                invoice_number TEXT NOT NULL,
                created_at     TEXT,
                patient_count  INTEGER DEFAULT 0,
                total          BIGINT  DEFAULT 0,
                patients_json  TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS personal_invoices (
                id            TEXT PRIMARY KEY,
                owner         TEXT NOT NULL,
                data_json     TEXT,
                created_at    TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS tasks (
                id          TEXT PRIMARY KEY,
                owner       TEXT NOT NULL,
                title       TEXT NOT NULL,
                description TEXT DEFAULT '',
                due_date    TEXT DEFAULT '',
                priority    TEXT DEFAULT 'normal',
                category    TEXT DEFAULT 'general',
                status      TEXT DEFAULT 'pendiente',
                reminder    TEXT DEFAULT '',
                created_at  TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS login_logs (
                id         TEXT PRIMARY KEY,
                user_id    TEXT NOT NULL,
                username   TEXT NOT NULL,
                name       TEXT NOT NULL,
                ip         TEXT DEFAULT '',
                device     TEXT DEFAULT '',
                status     TEXT DEFAULT 'exitoso',
                created_at TEXT
            )
        """)
    else:
        # SQLite
        db_execute("""
            CREATE TABLE IF NOT EXISTS users (
                id TEXT PRIMARY KEY, username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL, name TEXT NOT NULL,
                role TEXT DEFAULT 'user', active INTEGER DEFAULT 1,
                modules TEXT DEFAULT 'medical', created TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY, user_id TEXT NOT NULL, expires TEXT NOT NULL
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS medical_history (
                id TEXT PRIMARY KEY, owner TEXT NOT NULL,
                invoice_number TEXT NOT NULL, created_at TEXT,
                patient_count INTEGER DEFAULT 0, total INTEGER DEFAULT 0,
                patients_json TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS personal_invoices (
                id TEXT PRIMARY KEY, owner TEXT NOT NULL,
                data_json TEXT, created_at TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS tasks (
                id TEXT PRIMARY KEY, owner TEXT NOT NULL,
                title TEXT NOT NULL, description TEXT DEFAULT '',
                due_date TEXT DEFAULT '', priority TEXT DEFAULT 'normal',
                category TEXT DEFAULT 'general', status TEXT DEFAULT 'pendiente',
                reminder TEXT DEFAULT '', created_at TEXT
            )
        """)
        db_execute("""
            CREATE TABLE IF NOT EXISTS login_logs (
                id TEXT PRIMARY KEY, user_id TEXT NOT NULL,
                username TEXT NOT NULL, name TEXT NOT NULL,
                ip TEXT DEFAULT '', device TEXT DEFAULT '',
                status TEXT DEFAULT 'exitoso', created_at TEXT
            )
        """)

    # Admin por defecto
    admin = row_to_dict(db_execute(
        "SELECT id FROM users WHERE username = %s", ("admin",), fetch="one"
    ))
    if not admin:
        db_execute(
            "INSERT INTO users (id,username,password,name,role,active,modules,created) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (str(uuid.uuid4()), "admin", hash_password("admin123"),
             "Administrador", "admin", True, "medical,personal,tasks",
             datetime.now().isoformat())
        )
        print("âœ… Usuario admin creado (admin/admin123)")

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  AUTH HELPERS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def hash_password(pwd): return hashlib.sha256(pwd.encode()).hexdigest()

init_db()

def create_session(user_id):
    token   = secrets.token_hex(32)
    expires = (datetime.now() + timedelta(hours=8)).isoformat()
    db_execute("INSERT INTO sessions (token,user_id,expires) VALUES (%s,%s,%s)",
               (token, user_id, expires))
    return token

def get_session_user(token):
    if not token: return None
    sess = row_to_dict(db_execute(
        "SELECT * FROM sessions WHERE token=%s", (token,), fetch="one"
    ))
    if not sess: return None
    if datetime.fromisoformat(sess["expires"]) < datetime.now():
        db_execute("DELETE FROM sessions WHERE token=%s", (token,))
        return None
    user = row_to_dict(db_execute(
        "SELECT * FROM users WHERE id=%s", (sess["user_id"],), fetch="one"
    ))
    return user

def delete_session(token):
    db_execute("DELETE FROM sessions WHERE token=%s", (token,))

def require_auth(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        token = request.headers.get("Authorization","").replace("Bearer ","")
        user  = get_session_user(token)
        active = user.get("active") if user else False
        if isinstance(active, int): active = bool(active)
        if not user or not active:
            return jsonify(error="No autorizado"), 401
        g.user = user
        return f(*args, **kwargs)
    return wrapper

def require_admin(f):
    @wraps(f)
    def wrapper(*args, **kwargs):
        token = request.headers.get("Authorization","").replace("Bearer ","")
        user  = get_session_user(token)
        if not user or user.get("role") != "admin":
            return jsonify(error="Solo administradores"), 403
        g.user = user
        return f(*args, **kwargs)
    return wrapper

def require_module(module):
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            token = request.headers.get("Authorization","").replace("Bearer ","")
            user  = get_session_user(token)
            active = user.get("active") if user else False
            if isinstance(active, int): active = bool(active)
            if not user or not active:
                return jsonify(error="No autorizado"), 401
            mods = user.get("modules","").split(",")
            if module not in mods and user.get("role") != "admin":
                return jsonify(error="Sin acceso a este mÃ³dulo"), 403
            g.user = user
            return f(*args, **kwargs)
        return wrapper
    return decorator

def user_modules(user):
    return user.get("modules","").split(",")

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  API AUTH
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def log_login(user_id, username, name, status):
    ip     = request.headers.get("X-Forwarded-For", request.remote_addr or "")
    if "," in ip: ip = ip.split(",")[0].strip()
    ua     = request.headers.get("User-Agent","")
    if "Mobile" in ua or "Android" in ua or "iPhone" in ua:
        device = "ğŸ“± MÃ³vil"
    elif "Tablet" in ua or "iPad" in ua:
        device = "ğŸ“Ÿ Tablet"
    else:
        device = "ğŸ’» Escritorio"
    db_execute(
        "INSERT INTO login_logs (id,user_id,username,name,ip,device,status,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
        (str(uuid.uuid4()), user_id, username, name, ip, device, status, datetime.now().isoformat())
    )

@app.route("/api/auth/login", methods=["POST"])
def login():
    data = request.get_json(force=True) or {}
    username = data.get("username","").strip()
    # Check user exists first (for failed login logging)
    user = row_to_dict(db_execute(
        "SELECT * FROM users WHERE username=%s AND password=%s",
        (username, hash_password(data.get("password",""))),
        fetch="one"
    ))
    if not user:
        # Try to find user by username to log the attempt
        u2 = row_to_dict(db_execute("SELECT * FROM users WHERE username=%s",(username,),fetch="one"))
        if u2: log_login(u2["id"], u2["username"], u2["name"], "fallido")
        return jsonify(error="Usuario o contraseÃ±a incorrectos"), 401
    active = user.get("active")
    if isinstance(active, int): active = bool(active)
    if not active:
        log_login(user["id"], user["username"], user["name"], "bloqueado")
        return jsonify(error="Usuario bloqueado"), 403
    token = create_session(user["id"])
    mods  = user_modules(user)
    log_login(user["id"], user["username"], user["name"], "exitoso")
    return jsonify(token=token, user={
        "id":user["id"],"name":user["name"],"username":user["username"],
        "role":user["role"],"modules":mods
    })

@app.route("/api/auth/logout", methods=["POST"])
def logout():
    token = request.headers.get("Authorization","").replace("Bearer ","")
    user  = get_session_user(token)
    if user:
        log_login(user["id"], user["username"], user["name"], "cierre")
    delete_session(token)
    return jsonify(ok=True)

@app.route("/api/auth/me", methods=["GET"])
@require_auth
def me():
    u = g.user
    return jsonify(id=u["id"],name=u["name"],username=u["username"],
                   role=u["role"],modules=user_modules(u))

@app.route("/api/admin/login-logs", methods=["GET"])
@require_admin
def admin_login_logs():
    limit = int(request.args.get("limit", 100))
    rows = rows_to_list(db_execute(
        "SELECT * FROM login_logs ORDER BY created_at DESC LIMIT %s",
        (limit,), fetch="all"
    ))
    return jsonify(logs=rows)

@app.route("/api/auth/my-logs", methods=["GET"])
@require_auth
def my_login_logs():
    uid = g.user["id"]
    rows = rows_to_list(db_execute(
        "SELECT * FROM login_logs WHERE user_id=%s ORDER BY created_at DESC LIMIT 50",
        (uid,), fetch="all"
    ))
    return jsonify(logs=rows)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  API ADMIN
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.route("/api/admin/users", methods=["GET"])
@require_admin
def admin_list_users():
    users = rows_to_list(db_execute(
        "SELECT id,username,name,role,active,modules,created FROM users ORDER BY created",
        fetch="all"
    ))
    for u in users:
        u["modules"] = u.get("modules","").split(",")
        active = u.get("active")
        u["active"] = bool(active) if isinstance(active,int) else active
    return jsonify(users=users)

@app.route("/api/admin/users", methods=["POST"])
@require_admin
def admin_create_user():
    data = request.get_json(force=True) or {}
    username = data.get("username","").strip()
    password = data.get("password","").strip()
    name     = data.get("name","").strip()
    role     = data.get("role","user")
    modules  = ",".join(data.get("modules",["medical"]))
    if not username or not password or not name:
        return jsonify(error="Campos requeridos"), 400
    existing = db_execute("SELECT id FROM users WHERE username=%s",(username,),fetch="one")
    if existing: return jsonify(error="Usuario ya existe"), 409
    uid = str(uuid.uuid4())
    db_execute(
        "INSERT INTO users (id,username,password,name,role,active,modules,created) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
        (uid, username, hash_password(password), name, role, True, modules, datetime.now().isoformat())
    )
    return jsonify(id=uid,username=username,name=name,role=role,active=True,
                   modules=modules.split(","),created=datetime.now().isoformat()), 201

@app.route("/api/admin/users/<uid>", methods=["PUT"])
@require_admin
def admin_update_user(uid):
    data = request.get_json(force=True) or {}
    user = row_to_dict(db_execute("SELECT * FROM users WHERE id=%s",(uid,),fetch="one"))
    if not user: return jsonify(error="No encontrado"), 404
    name    = data.get("name", user["name"])
    role    = data.get("role", user["role"])
    active  = data.get("active", user["active"])
    modules = ",".join(data.get("modules", user.get("modules","").split(",")))
    if isinstance(active, int): active = bool(active)
    if "password" in data and data["password"].strip():
        db_execute("UPDATE users SET name=%s,role=%s,active=%s,modules=%s,password=%s WHERE id=%s",
                   (name,role,active,modules,hash_password(data["password"].strip()),uid))
    else:
        db_execute("UPDATE users SET name=%s,role=%s,active=%s,modules=%s WHERE id=%s",
                   (name,role,active,modules,uid))
    return jsonify(id=uid,name=name,role=role,active=active,modules=modules.split(","))

@app.route("/api/admin/users/<uid>", methods=["DELETE"])
@require_admin
def admin_delete_user(uid):
    user = row_to_dict(db_execute("SELECT * FROM users WHERE id=%s",(uid,),fetch="one"))
    if not user: return jsonify(error="No encontrado"), 404
    if user.get("username") == "admin": return jsonify(error="No puedes eliminar el admin principal"), 400
    db_execute("DELETE FROM users WHERE id=%s",(uid,))
    return "", 204

@app.route("/api/admin/users/<uid>/toggle", methods=["POST"])
@require_admin
def admin_toggle_user(uid):
    user = row_to_dict(db_execute("SELECT * FROM users WHERE id=%s",(uid,),fetch="one"))
    if not user: return jsonify(error="No encontrado"), 404
    if user.get("username") == "admin": return jsonify(error="No puedes bloquear al admin"), 400
    active = user.get("active")
    if isinstance(active, int): active = bool(active)
    new_active = not active
    db_execute("UPDATE users SET active=%s WHERE id=%s",(new_active,uid))
    return jsonify(active=new_active)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  CONSTANTES MÃ‰DICAS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

DOCTOR  = "DR. FRANCISCO ENRIQUE CABRERA PORTIELES"
SPEC    = "NEUROFISIOLOGO CLINICO"
LICENSE = "RM0307 - CC 1047488543"

patients_store = {}  # RAM: { user_id: [patients] }

def get_patients(uid): return patients_store.setdefault(uid, [])
def auto_price(idx): return 100_000 if idx < 20 else 70_000
def clean(name): return Path(name).stem.replace("_"," ").title()
def fmt_money(v): return f"${int(v):,}".replace(",",".")

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  GENERADORES DE DOCUMENTOS MÃ‰DICOS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def docx_invoice(number, patients):
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2); s.bottom_margin=Cm(2)
        s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    style=doc.styles['Normal']; style.font.name='Arial'; style.font.size=Pt(11)
    hdr=doc.add_paragraph(); hdr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=hdr.add_run(DOCTOR+"\n"); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(0,51,102)
    r2=hdr.add_run(SPEC+"\n"); r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(0,51,102)
    r3=hdr.add_run(LICENSE+"\n\n"); r3.italic=True; r3.font.size=Pt(10); r3.font.color.rgb=RGBColor(0,51,102)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rc=p.add_run(f"SE REALIZÃ“ INFORME Y PROCESAMIENTO DE LA CANTIDAD DE ESTUDIOS: {len(patients)}\nESTUDIOS DE POLISOMNOGRAFÃA\n\n")
    rc.bold=True; rc.font.size=Pt(11); rc.font.color.rgb=RGBColor(0,51,102)
    doc.add_paragraph(f"FACTURA NÂ°: {number}",style='Heading 1').runs[0].bold=True
    doc.add_paragraph(f"Fecha: {datetime.now():%d/%m/%Y %H:%M}"); doc.add_paragraph()
    tbl=doc.add_table(rows=1,cols=3); tbl.style='Light List Accent 1'
    tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hc=tbl.rows[0].cells; hc[0].text="No."; hc[1].text="PACIENTE"; hc[2].text="VALOR"
    for i,p in enumerate(patients,1):
        row=tbl.add_row().cells; row[0].text=str(i); row[1].text=p['name']; row[2].text=fmt_money(p['price'])
    doc.add_paragraph(); total=sum(p['price'] for p in patients)
    tp=doc.add_paragraph(); tp.paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
    rt=tp.add_run(f"TOTAL: {fmt_money(total)}"); rt.bold=True; rt.font.size=Pt(12)
    fn=TEMP_DIR/f"Factura_{number}.docx"; doc.save(fn); return fn

def generate_pdf(number, patients):
    path=TEMP_DIR/f"Factura_{number}.pdf"
    c=canvas.Canvas(str(path),pagesize=letter); w,h=letter
    def tbl_header(y):
        c.setFont("Helvetica-Bold",10)
        c.drawString(72,y,"No."); c.drawString(120,y,"PACIENTE"); c.drawRightString(w-72,y,"VALOR")
        c.line(72,y-5,w-72,y-5)
    c.setFont("Helvetica-Bold",14); c.setFillColorRGB(0,51/255,102/255)
    c.drawCentredString(w/2,h-50,DOCTOR); c.setFont("Helvetica",12)
    c.drawCentredString(w/2,h-70,SPEC); c.drawCentredString(w/2,h-85,LICENSE)
    c.setFont("Helvetica-Bold",10)
    c.drawCentredString(w/2,h-105,f"SE REALIZÃ“ INFORME Y PROCESAMIENTO DE LA CANTIDAD DE ESTUDIOS: {len(patients)}")
    c.drawCentredString(w/2,h-120,"ESTUDIOS DE POLISOMNOGRAFÃA")
    c.setFont("Helvetica-Bold",12); c.drawString(72,h-150,f"FACTURA NÂ°: {number}")
    c.setFont("Helvetica",11); c.drawString(72,h-170,f"Fecha: {datetime.now():%d/%m/%Y %H:%M}")
    rh=18; mr=int((h-260)/rh); y=h-225; tbl_header(h-200)
    for i,p in enumerate(patients,1):
        if (i-1) and (i-1)%mr==0: c.showPage(); tbl_header(h-50); y=h-80
        c.setFont("Helvetica",10); c.drawString(72,y,str(i))
        c.drawString(120,y,p['name']); c.drawRightString(w-72,y,fmt_money(p['price'])); y-=rh
    sub=sum(p['price'] for p in patients)
    if y<100: c.showPage(); tbl_header(h-50); y=h-80
    c.setFont("Helvetica-Bold",11); c.drawRightString(w-72,y-20,f"SUBTOTAL: {fmt_money(sub)}")
    c.setFont("Helvetica-Bold",12); c.drawRightString(w-72,y-40,f"TOTAL:    {fmt_money(sub)}")
    c.save(); return path

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  API MÃ‰DICA
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.route("/api/medical/patients", methods=["GET","POST"])
@require_module("medical")
def medical_patients():
    uid=g.user["id"]; pts=get_patients(uid)
    if request.method=="POST" and "files" in request.files:
        new=[]
        for f in request.files.getlist("files"):
            if f.filename.lower().endswith((".doc",".docx",".pdf")):
                idx=len(pts)+len(new)
                new.append({"id":idx+1,"name":clean(f.filename),"price":auto_price(idx)})
        pts.extend(new); return jsonify(success=True,patients=new)
    if request.method=="POST":
        data=request.get_json(force=True) or {}
        name=data.get("name","").strip()
        if not name: return jsonify(error="Nombre requerido"),400
        idx=len(pts); price=int(data.get("price") or auto_price(idx))
        if idx<20: price=100_000
        p={"id":idx+1,"name":name,"price":price}; pts.append(p)
        return jsonify(p),201
    sub=sum(p["price"] for p in pts)
    return jsonify(patients=pts,count=len(pts),subtotal=sub)

@app.route("/api/medical/patients/<int:pid>", methods=["PUT","DELETE"])
@require_module("medical")
def medical_one(pid):
    uid=g.user["id"]; pts=get_patients(uid)
    p=next((x for x in pts if x["id"]==pid),None)
    if not p: return "",404
    if request.method=="DELETE":
        pts.remove(p)
        for i,obj in enumerate(pts,1): obj["id"]=i
        return "",204
    data=request.get_json(force=True) or {}
    p["name"]=data.get("name",p["name"]).strip()
    if "price" in data:
        try: p["price"]=int(data["price"])
        except: return jsonify(error="Precio invÃ¡lido"),400
    else: p["price"]=auto_price(p["id"]-1)
    return jsonify(p)

@app.route("/api/medical/clear", methods=["DELETE"])
@require_module("medical")
def medical_clear():
    get_patients(g.user["id"]).clear(); return "",204

@app.route("/api/medical/invoice/<fmt>", methods=["POST"])
@require_module("medical")
def medical_invoice(fmt):
    uid=g.user["id"]; pts=get_patients(uid)
    if not pts: return jsonify(error="No hay pacientes"),400
    data=request.get_json(force=True) or {}
    num=data.get("invoice_number",f"FAC-{datetime.now():%Y%m%d%H%M%S}")
    # Guardar en historial
    total=sum(p["price"] for p in pts)
    db_execute(
        "INSERT INTO medical_history (id,owner,invoice_number,created_at,patient_count,total,patients_json) VALUES (%s,%s,%s,%s,%s,%s,%s)",
        (str(uuid.uuid4()), uid, num, datetime.now().isoformat(), len(pts), total, json.dumps(pts))
    )
    path=docx_invoice(num,pts) if fmt=="word" else generate_pdf(num,pts)
    return send_file(path,as_attachment=True)

@app.route("/api/medical/history", methods=["GET"])
@require_module("medical")
def medical_history():
    uid=g.user["id"]
    rows=rows_to_list(db_execute(
        "SELECT id,invoice_number,created_at,patient_count,total FROM medical_history WHERE owner=%s ORDER BY created_at DESC LIMIT 50",
        (uid,),fetch="all"
    ))
    return jsonify(history=rows)

@app.route("/api/medical/history/<hid>/download/<fmt>", methods=["GET"])
@require_module("medical")
def medical_history_download(hid,fmt):
    uid=g.user["id"]
    row=row_to_dict(db_execute(
        "SELECT * FROM medical_history WHERE id=%s AND owner=%s",(hid,uid),fetch="one"
    ))
    if not row: return jsonify(error="No encontrado"),404
    pts=json.loads(row["patients_json"])
    num=row["invoice_number"]
    path=docx_invoice(num,pts) if fmt=="word" else generate_pdf(num,pts)
    return send_file(path,as_attachment=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  API PERSONAL â€” generadores PDF/DOCX
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def generate_personal_pdf(inv):
    num=inv["number"]; path=TEMP_DIR/f"FacturaPersonal_{num}.pdf"
    c=canvas.Canvas(str(path),pagesize=A4); w,h=A4
    c.setFillColorRGB(.05,.08,.18); c.rect(0,h-120,w,120,fill=1,stroke=0)
    c.setFillColorRGB(.24,.52,1.0); c.rect(0,h-124,w,4,fill=1,stroke=0)
    c.setFont("Helvetica-Bold",20); c.setFillColorRGB(1,1,1)
    c.drawString(40,h-55,inv.get("issuer_name","").upper())
    c.setFont("Helvetica",11); c.setFillColorRGB(.7,.8,1.0)
    c.drawString(40,h-75,inv.get("issuer_email",""))
    c.drawString(40,h-90,inv.get("issuer_phone",""))
    c.drawString(40,h-105,inv.get("issuer_address",""))
    c.setFont("Helvetica-Bold",28); c.setFillColorRGB(1,1,1)
    c.drawRightString(w-40,h-60,"FACTURA")
    c.setFont("Helvetica",13); c.setFillColorRGB(.7,.8,1.0)
    c.drawRightString(w-40,h-80,f"NÂ° {num}")
    c.drawRightString(w-40,h-97,f"Fecha: {inv.get('date',datetime.now().strftime('%d/%m/%Y'))}")
    c.setFillColorRGB(.94,.96,1.0); c.rect(40,h-210,w-80,75,fill=1,stroke=0)
    c.setFont("Helvetica-Bold",9); c.setFillColorRGB(.24,.52,1.0)
    c.drawString(55,h-148,"FACTURADO A")
    c.setFont("Helvetica-Bold",13); c.setFillColorRGB(.05,.08,.18)
    c.drawString(55,h-165,inv.get("client_name",""))
    c.setFont("Helvetica",10); c.setFillColorRGB(.3,.35,.5)
    c.drawString(55,h-180,inv.get("client_company",""))
    c.drawString(55,h-193,inv.get("client_nit",""))
    c.drawString(55,h-206,inv.get("client_email",""))
    sc={"pagada":(0.13,0.7,0.4),"pendiente":(0.95,0.6,0.1),"vencida":(0.9,0.2,0.2)}.get(inv.get("status","pendiente"),(0.5,0.5,0.5))
    c.setFillColorRGB(*sc); c.roundRect(w-160,h-168,110,22,5,fill=1,stroke=0)
    c.setFont("Helvetica-Bold",10); c.setFillColorRGB(1,1,1)
    c.drawCentredString(w-105,h-153,inv.get("status","pendiente").upper())
    ty=h-255; c.setFillColorRGB(.05,.08,.18); c.rect(40,ty,w-80,24,fill=1,stroke=0)
    c.setFont("Helvetica-Bold",9); c.setFillColorRGB(1,1,1)
    c.drawString(55,ty+8,"DESCRIPCIÃ“N"); c.drawRightString(w/2+20,ty+8,"CANT.")
    c.drawRightString(w/2+120,ty+8,"V. UNITARIO"); c.drawRightString(w-50,ty+8,"TOTAL")
    ry=ty-5; rh=22; sub=0
    for i,it in enumerate(inv.get("items",[])):
        bg=(.97,.98,1.) if i%2==0 else (1.,1.,1.)
        c.setFillColorRGB(*bg); c.rect(40,ry-rh+4,w-80,rh,fill=1,stroke=0)
        qty=float(it.get("qty",1)); uv=float(it.get("unit_value",0)); tot=qty*uv; sub+=tot
        c.setFont("Helvetica",9); c.setFillColorRGB(.1,.12,.25)
        c.drawString(55,ry-4,it.get("description","")[:65])
        c.drawRightString(w/2+20,ry-4,str(int(qty) if qty==int(qty) else qty))
        c.drawRightString(w/2+120,ry-4,f"${uv:,.0f}".replace(",","."))
        c.drawRightString(w-50,ry-4,f"${tot:,.0f}".replace(",",".")); ry-=rh
    tax=float(inv.get("tax",0)); taxv=sub*tax/100; total=sub+taxv
    toty=ry-15; c.setFillColorRGB(.94,.96,1.0); c.rect(w/2+20,toty-50,w/2-60,65,fill=1,stroke=0)
    c.setFont("Helvetica",10); c.setFillColorRGB(.3,.35,.5)
    c.drawString(w/2+35,toty+5,"Subtotal:"); c.drawRightString(w-55,toty+5,f"${sub:,.0f}".replace(",","."))
    if tax:
        c.drawString(w/2+35,toty-13,f"IVA ({tax:.0f}%):"); c.drawRightString(w-55,toty-13,f"${taxv:,.0f}".replace(",","."))
    c.setFillColorRGB(.05,.08,.18); c.rect(w/2+20,toty-50,w/2-60,22,fill=1,stroke=0)
    c.setFont("Helvetica-Bold",12); c.setFillColorRGB(1,1,1)
    c.drawString(w/2+35,toty-41,"TOTAL:"); c.drawRightString(w-55,toty-41,f"${total:,.0f}".replace(",","."))
    notes=inv.get("notes","")
    if notes:
        c.setFont("Helvetica-Bold",9); c.setFillColorRGB(.3,.35,.5)
        c.drawString(40,toty-70,"NOTAS:"); c.setFont("Helvetica",9)
        c.drawString(40,toty-83,notes[:100])
    c.setFillColorRGB(.05,.08,.18); c.rect(0,0,w,40,fill=1,stroke=0)
    c.setFont("Helvetica",8); c.setFillColorRGB(.5,.6,.8)
    c.drawCentredString(w/2,25,"Generado con Facturador FL Â· "+datetime.now().strftime("%d/%m/%Y %H:%M"))
    c.save(); return path

def generate_personal_docx(inv):
    num=inv["number"]; doc=Document()
    for s in doc.sections:
        s.top_margin=Cm(2);s.bottom_margin=Cm(2);s.left_margin=Cm(2.5);s.right_margin=Cm(2.5)
    hdr=doc.add_paragraph(); hdr.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=hdr.add_run(inv.get("issuer_name","").upper()+"\n"); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=RGBColor(13,20,46)
    r2=hdr.add_run(f"{inv.get('issuer_email','')}  |  {inv.get('issuer_phone','')}\n{inv.get('issuer_address','')}\n")
    r2.font.size=Pt(9); r2.font.color.rgb=RGBColor(100,120,160)
    pt=doc.add_paragraph(); pt.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rt=pt.add_run(f"FACTURA NÂ° {num}"); rt.bold=True; rt.font.size=Pt(22); rt.font.color.rgb=RGBColor(13,20,46)
    doc.add_paragraph(f"Fecha: {inv.get('date','')}  |  Vence: {inv.get('due_date','')}")
    doc.add_paragraph(); doc.add_paragraph("FACTURADO A:").runs[0].bold=True
    doc.add_paragraph(f"{inv.get('client_name','')}  â€“  {inv.get('client_company','')}")
    doc.add_paragraph(f"NIT/CC: {inv.get('client_nit','')}  |  Email: {inv.get('client_email','')}"); doc.add_paragraph()
    tbl=doc.add_table(rows=1,cols=4); tbl.style='Light List Accent 1'
    hc=tbl.rows[0].cells; hc[0].text="DESCRIPCIÃ“N"; hc[1].text="CANT."; hc[2].text="V.UNITARIO"; hc[3].text="TOTAL"
    sub=0
    for it in inv.get("items",[]):
        qty=float(it.get("qty",1)); uv=float(it.get("unit_value",0)); tot=qty*uv; sub+=tot
        row=tbl.add_row().cells; row[0].text=it.get("description","")
        row[1].text=str(int(qty) if qty==int(qty) else qty)
        row[2].text=f"${uv:,.0f}".replace(",","."); row[3].text=f"${tot:,.0f}".replace(",",".")
    doc.add_paragraph(); tax=float(inv.get("tax",0)); taxv=sub*tax/100; total=sub+taxv
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rt=tp.add_run(f"Subtotal: ${sub:,.0f}\n"+(f"IVA: ${taxv:,.0f}\n" if tax else "")+f"TOTAL: ${total:,.0f}")
    rt.bold=True; rt.font.size=Pt(12)
    if inv.get("notes"): doc.add_paragraph(); doc.add_paragraph(f"Notas: {inv['notes']}")
    fn=TEMP_DIR/f"FacturaPersonal_{num}.docx"; doc.save(fn); return fn

# â”€â”€ Endpoints personales â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

@app.route("/api/personal/invoices", methods=["GET"])
@require_module("personal")
def personal_list():
    uid=g.user["id"]
    rows=rows_to_list(db_execute(
        "SELECT id,data_json,created_at FROM personal_invoices WHERE owner=%s ORDER BY created_at DESC",
        (uid,),fetch="all"
    ))
    invs=[json.loads(r["data_json"]) for r in rows]
    return jsonify(invoices=invs)

@app.route("/api/personal/invoices", methods=["POST"])
@require_module("personal")
def personal_create():
    data=request.get_json(force=True) or {}; uid=g.user["id"]
    inv={
        "id":str(uuid.uuid4()),"owner":uid,
        "number":data.get("number",f"INV-{datetime.now():%Y%m%d%H%M%S}"),
        "date":data.get("date",datetime.now().strftime("%d/%m/%Y")),
        "due_date":data.get("due_date",""),"status":data.get("status","pendiente"),
        "issuer_name":data.get("issuer_name",""),"issuer_email":data.get("issuer_email",""),
        "issuer_phone":data.get("issuer_phone",""),"issuer_address":data.get("issuer_address",""),
        "client_name":data.get("client_name",""),"client_company":data.get("client_company",""),
        "client_nit":data.get("client_nit",""),"client_email":data.get("client_email",""),
        "items":data.get("items",[]),"tax":data.get("tax",0),"notes":data.get("notes",""),
        "created":datetime.now().isoformat()
    }
    db_execute(
        "INSERT INTO personal_invoices (id,owner,data_json,created_at) VALUES (%s,%s,%s,%s)",
        (inv["id"],uid,json.dumps(inv),datetime.now().isoformat())
    )
    return jsonify(inv),201

@app.route("/api/personal/invoices/<iid>", methods=["PUT"])
@require_module("personal")
def personal_update(iid):
    uid=g.user["id"]
    row=row_to_dict(db_execute("SELECT * FROM personal_invoices WHERE id=%s AND owner=%s",(iid,uid),fetch="one"))
    if not row: return jsonify(error="No encontrado"),404
    inv=json.loads(row["data_json"])
    data=request.get_json(force=True) or {}
    for f in ["number","date","due_date","status","issuer_name","issuer_email","issuer_phone",
              "issuer_address","client_name","client_company","client_nit","client_email","items","tax","notes"]:
        if f in data: inv[f]=data[f]
    db_execute("UPDATE personal_invoices SET data_json=%s WHERE id=%s",(json.dumps(inv),iid))
    return jsonify(inv)

@app.route("/api/personal/invoices/<iid>", methods=["DELETE"])
@require_module("personal")
def personal_delete(iid):
    uid=g.user["id"]
    row=row_to_dict(db_execute("SELECT id FROM personal_invoices WHERE id=%s AND owner=%s",(iid,uid),fetch="one"))
    if not row: return jsonify(error="No encontrado"),404
    db_execute("DELETE FROM personal_invoices WHERE id=%s",(iid,))
    return "",204

@app.route("/api/personal/invoices/<iid>/download/<fmt>", methods=["GET"])
@require_module("personal")
def personal_download(iid,fmt):
    uid=g.user["id"]
    row=row_to_dict(db_execute("SELECT * FROM personal_invoices WHERE id=%s AND owner=%s",(iid,uid),fetch="one"))
    if not row: return jsonify(error="No encontrado"),404
    inv=json.loads(row["data_json"])
    path=generate_personal_docx(inv) if fmt=="word" else generate_personal_pdf(inv)
    return send_file(path,as_attachment=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  API TAREAS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.route("/api/tasks", methods=["GET"])
@require_module("tasks")
def tasks_list():
    uid=g.user["id"]
    status_filter=request.args.get("status","")
    if status_filter:
        rows=rows_to_list(db_execute(
            "SELECT * FROM tasks WHERE owner=%s AND status=%s ORDER BY due_date,created_at DESC",
            (uid,status_filter),fetch="all"
        ))
    else:
        rows=rows_to_list(db_execute(
            "SELECT * FROM tasks WHERE owner=%s ORDER BY due_date,created_at DESC",
            (uid,),fetch="all"
        ))
    return jsonify(tasks=rows)

@app.route("/api/tasks", methods=["POST"])
@require_module("tasks")
def tasks_create():
    data=request.get_json(force=True) or {}; uid=g.user["id"]
    title=data.get("title","").strip()
    if not title: return jsonify(error="TÃ­tulo requerido"),400
    tid=str(uuid.uuid4())
    db_execute(
        "INSERT INTO tasks (id,owner,title,description,due_date,priority,category,status,reminder,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
        (tid,uid,title,
         data.get("description",""),
         data.get("due_date",""),
         data.get("priority","normal"),
         data.get("category","general"),
         data.get("status","pendiente"),
         data.get("reminder",""),
         datetime.now().isoformat())
    )
    task=row_to_dict(db_execute("SELECT * FROM tasks WHERE id=%s",(tid,),fetch="one"))
    return jsonify(task),201

@app.route("/api/tasks/<tid>", methods=["PUT"])
@require_module("tasks")
def tasks_update(tid):
    uid=g.user["id"]
    task=row_to_dict(db_execute("SELECT * FROM tasks WHERE id=%s AND owner=%s",(tid,uid),fetch="one"))
    if not task: return jsonify(error="No encontrado"),404
    data=request.get_json(force=True) or {}
    title      =data.get("title",      task["title"])
    description=data.get("description",task["description"])
    due_date   =data.get("due_date",   task["due_date"])
    priority   =data.get("priority",   task["priority"])
    category   =data.get("category",   task["category"])
    status     =data.get("status",     task["status"])
    reminder   =data.get("reminder",   task["reminder"])
    db_execute(
        "UPDATE tasks SET title=%s,description=%s,due_date=%s,priority=%s,category=%s,status=%s,reminder=%s WHERE id=%s",
        (title,description,due_date,priority,category,status,reminder,tid)
    )
    task=row_to_dict(db_execute("SELECT * FROM tasks WHERE id=%s",(tid,),fetch="one"))
    return jsonify(task)

@app.route("/api/tasks/<tid>", methods=["DELETE"])
@require_module("tasks")
def tasks_delete(tid):
    uid=g.user["id"]
    task=row_to_dict(db_execute("SELECT id FROM tasks WHERE id=%s AND owner=%s",(tid,uid),fetch="one"))
    if not task: return jsonify(error="No encontrado"),404
    db_execute("DELETE FROM tasks WHERE id=%s",(tid,))
    return "",204

@app.route("/api/tasks/<tid>/complete", methods=["POST"])
@require_module("tasks")
def tasks_complete(tid):
    uid=g.user["id"]
    task=row_to_dict(db_execute("SELECT * FROM tasks WHERE id=%s AND owner=%s",(tid,uid),fetch="one"))
    if not task: return jsonify(error="No encontrado"),404
    new_status="completada" if task["status"]!="completada" else "pendiente"
    db_execute("UPDATE tasks SET status=%s WHERE id=%s",(new_status,tid))
    return jsonify(status=new_status)

@app.route("/api/tasks/reminders", methods=["GET"])
@require_module("tasks")
def tasks_reminders():
    """Tareas con recordatorio en las prÃ³ximas 24h o vencidas"""
    uid=g.user["id"]
    now=datetime.now()
    tomorrow=(now+timedelta(hours=24)).strftime("%Y-%m-%d")
    today=now.strftime("%Y-%m-%d")
    rows=rows_to_list(db_execute(
        "SELECT * FROM tasks WHERE owner=%s AND status!='completada' AND due_date!='' AND due_date<=%s ORDER BY due_date",
        (uid,tomorrow),fetch="all"
    ))
    reminders=[]
    for t in rows:
        due=t.get("due_date","")
        if due:
            try:
                due_dt=datetime.strptime(due,"%Y-%m-%d")
                overdue=due_dt.date()<now.date()
                due_today=due_dt.date()==now.date()
                due_tomorrow=due_dt.date()==(now+timedelta(days=1)).date()
                t["alert_type"]="vencida" if overdue else ("hoy" if due_today else "maÃ±ana")
                reminders.append(t)
            except: pass
    return jsonify(reminders=reminders)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
#  RUTAS ESTÃTICAS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

@app.route("/")
def home(): return send_from_directory(app.static_folder,"login.html")

@app.route("/dashboard")
def dashboard(): return send_from_directory(app.static_folder,"dashboard.html")

@app.route("/medical")
def medical(): return send_from_directory(app.static_folder,"index.html")

@app.route("/personal")
def personal(): return send_from_directory(app.static_folder,"personal.html")

@app.route("/admin")
def admin_panel(): return send_from_directory(app.static_folder,"admin.html")

@app.route("/tasks")
def tasks_page(): return send_from_directory(app.static_folder,"tasks.html")

@app.route("/login-history")
def login_history_page(): return send_from_directory(app.static_folder,"login-history.html")

if __name__ == "__main__":
    host=os.environ.get("HOST","0.0.0.0")
    port=int(os.environ.get("PORT",5000))
    app.run(host=host,port=port,debug=True)
